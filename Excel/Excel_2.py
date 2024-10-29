import pandas as pd
from docx import Document

def df_to_word_table(df, document, nan_replacement=" "):
   
    df = df.fillna(nan_replacement)
    table = document.add_table(rows=1, cols=len(df.columns))

    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(df.columns):
        hdr_cells[i].text = str(column_name)
 
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

def print_all_sheets_data_to_word(excel_path, output_word_file, nan_replacement=""):
    
    all_sheets_df = pd.read_excel(excel_path, sheet_name=None)

    # Create a new Document
    document = Document()

    # Loop through each sheet and add its data as a table in the Word document
    for sheet_name, df in all_sheets_df.items():
        document.add_heading(f'Sheet: {sheet_name}', level=1)  # Add sheet name as heading
        df_to_word_table(df, document, nan_replacement)  # Add DataFrame as a table, replacing NaN
        document.add_page_break()  # Add page break between sheets

    # Save the document
    document.save(output_word_file)
    print(f"Data has been saved to {output_word_file}")

# Example usage
file_path = r'C:\Users\abhasker1\Downloads\EmilyTeohProject\ExcelData_Important\1.1 Anglo_PolicyData.xlsx'
output_word_file = 'output.docx'
print_all_sheets_data_to_word(file_path, output_word_file, nan_replacement=" ")
