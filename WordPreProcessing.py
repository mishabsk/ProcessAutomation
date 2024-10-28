from docx import Document


document=Document()
document.add_heading("Akshi's Pre-processing of Files")
p=document.add_paragraph("Let's learn on how to work with Word")

table_header=["Name", "Age", "School"]

data=[["Priya", "16", "DPS"],["Payal", "18", "GIS"], ["Kajaal", "18", "BIS"]]

table=document.add_table(rows=1, cols=3)

for i in range(3):
    table.rows[0].cells[i].text=table_header[i]

for name, age, school in data:
    cells=table.add_row().cells
    cells[0].text=name
    cells[1].text=age
    cells[2].text=school

document.save("Test.docx")
